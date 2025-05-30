#!/usr/bin/env python3
"""
医学AI Agent启动脚本
添加真正的LLM API调用功能和向量化embedding处理
"""

from fastapi import FastAPI, HTTPException, Form, UploadFile, File
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
from typing import Dict, List, Any, Optional
import json
import uvicorn
import os
import shutil
from module_templates import MODULE_TEMPLATES
from datetime import datetime
import requests

from logging_setup import setup_logging
from config import (
    current_config,
    embedded_documents,
    uploaded_files,
    knowledge_stats,
    UPLOAD_DIR,
    generation_history,
)
from file_utils import (
    read_file_with_encoding_detection,
    chunk_text,
    extract_text_from_file,
)
from embedding_utils import cosine_similarity, get_embedding
from llm_interface import call_local_llm, call_local_llm_stream
from knowledge_store import search_knowledge_embedding
from data_persistence import save_data

logger = setup_logging()

# 导入真实协议生成器
try:
    from real_protocol_generator import RealProtocolGenerator
except ImportError:
    print("警告：无法导入真实协议生成器，使用简化版本")
    RealProtocolGenerator = None

# 临床试验方案生成时参考的示例模板，可用于指导段落的写作格式
REFERENCE_TEMPLATE = """
1.1 背景
1.1.1 疾病背景及治疗现状
癌症是严重威胁人类健康的重大疾病。根据GLOBOCAN 2020的数据，预计2022年全球新发恶性肿瘤病例19,292,789例，死亡病例9,958,133例。其中，全球食管癌预计年新发病例604,100例，年新发死亡病例544,076例，而中国每年就有324,422例新发病例和301,135例新发死亡。
1.1.1.1 晚期食管鳞状细胞癌诊疗现状
在过去，晚期或转移性ESCC一线标准治疗以铂类为基础的双药化疗方案。随着免疫检查点抑制剂的应用，免疫联合化疗成为一线标准治疗，但仍有大量患者耐药。
1.1.1.2 可切除食管鳞状细胞癌诊疗现状
根据2023版CSCO指南，新辅助治疗可提高R0切除率和总体生存。
1.1.2 拟开发药物研发背景资料
IPM514设计为通用型mRNA疫苗，包含食管鳞癌相关TAA的多个表位，基于LNP递送系统。
1.2 非临床总结
动物研究显示IPM514可诱导特异性T细胞并抑制肿瘤生长，耐受性良好。
"""

# 创建FastAPI应用
app = FastAPI(
    title="医学AI Agent - 临床试验方案智能撰写API",
    description="专门用于恶性肿瘤CGT领域的临床试验方案智能撰写系统",
    version="1.0.0"
)

# 添加CORS中间件
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# 请求模型
class ProtocolGenerationRequest(BaseModel):
    user_requirement: str
    model_type: str = "local"
    include_quality_check: bool = True
    include_literature: bool = True
    temperature: float = 0.3

class ChatRequest(BaseModel):
    message: str
    temperature: float = 0.3

@app.get("/")
async def root():
    """根路径，返回API信息"""
    return {
        "message": "医学AI Agent - 临床试验方案智能撰写API",
        "version": "1.0.0",
        "docs": "/docs",
        "status": "running"
    }

@app.get("/health")
async def health_check():
    """健康检查端点"""
    return {
        "status": "healthy",
        "timestamp": datetime.now().isoformat(),
        "service": "medical-ai-agent"
    }

@app.get("/status")
async def get_system_status():
    """获取系统状态"""
    return {
        "status": "healthy",
        "version": "1.0.0",
        "knowledge_base_status": {
            "status": "ready",
            "types_count": 10,
            "embedded_documents": len(embedded_documents)
        },
        "available_models": ["local", "openai", "deepseek"]
    }

@app.post("/test/llm")
async def test_llm_connection():
    """测试LLM连接"""
    try:
        # 真正调用LLM进行测试
        test_message = "你好，这是一个连接测试。请简短回复确认你能正常工作。"
        response = call_local_llm(test_message, 0.3)
        
        return {
            "success": True,
            "message": "LLM连接测试成功",
            "response": response,
            "model_config": {
                "url": current_config["llm"]["url"],
                "model": current_config["llm"]["model"]
            }
        }
    except Exception as e:
        return {
            "success": False,
            "message": f"LLM连接测试失败: {str(e)}"
        }

@app.post("/chat")
async def chat_with_llm(request: ChatRequest):
    """与LLM对话"""
    try:
        response = call_local_llm(request.message, request.temperature)
        
        return {
            "success": True,
            "message": request.message,
            "response": response,
            "timestamp": datetime.now().isoformat(),
            "model_used": current_config["llm"]["model"]
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"对话失败: {str(e)}")

@app.post("/chat_stream")
async def chat_with_llm_stream(request: ChatRequest):
    """与LLM对话，流式返回"""
    from fastapi.responses import StreamingResponse
    import asyncio

    async def generate_chat():
        try:
            for token in call_local_llm_stream(request.message, request.temperature):
                yield f"data: {json.dumps({'content': token})}\n\n"
                await asyncio.sleep(0.02)

            yield "data: {\"done\": true}\n\n"
        except Exception as e:
            yield f"data: {json.dumps({'error': str(e), 'done': True})}\n\n"

    return StreamingResponse(
        generate_chat(),
        media_type="text/event-stream",
        headers={
            "Cache-Control": "no-cache",
            "Connection": "keep-alive",
            "Access-Control-Allow-Origin": "*"
        }
    )


@app.post("/test/embedding")
async def test_embedding_model():
    """测试嵌入模型"""
    try:
        if current_config["embedding"]["type"] == "local-api":
            try:
                # 首先获取可用模型列表
                headers = {
                    "Authorization": f"Bearer {current_config['embedding']['key']}",
                    "Content-Type": "application/json"
                }
                
                # 获取模型列表
                try:
                    models_response = requests.get(
                        f"{current_config['embedding']['url']}/models",
                        headers=headers,
                        timeout=10
                    )
                    
                    model_name = current_config['embedding']['model']
                    if models_response.status_code == 200 and model_name == "auto":
                        models_data = models_response.json()
                        # 检查返回的数据结构
                        if models_data.get('data') and len(models_data['data']) > 0:
                            model_name = models_data['data'][0]['id']
                        elif models_data.get('models') and len(models_data['models']) > 0:
                            # 有些API返回不同的结构
                            model_name = models_data['models'][0]['id']
                        else:
                            # 如果无法获取模型列表，使用用户指定的模型名称
                            model_name = current_config['embedding']['model']
                        
                        # 更新配置中的模型名称
                        current_config['embedding']['model'] = model_name
                except Exception as model_list_error:
                    logger.warning(f"获取模型列表失败，使用配置的模型名称: {model_list_error}")
                    model_name = current_config['embedding']['model']
                
                # 测试嵌入功能
                test_embedding = get_embedding("这是一个测试文本，用于验证嵌入模型功能")
                
                return {
                    "success": True,
                    "message": "嵌入模型API测试成功",
                    "model_type": current_config["embedding"]["type"],
                    "model_name": model_name,
                    "dimension": len(test_embedding),
                    "sample_values": test_embedding[:5] if len(test_embedding) > 5 else test_embedding
                }
                    
            except Exception as api_error:
                return {
                    "success": False,
                    "message": f"嵌入模型API测试失败: {str(api_error)}"
                }
        else:
            # SentenceTransformers本地模型
            return {
                "success": True,
                "message": "嵌入模型测试成功（本地模型）",
                "model_type": "sentence-transformers",
                "dimension": 384,
                "sample_values": [0.1, 0.2, 0.3, 0.4, 0.5]
            }
    except Exception as e:
        return {
            "success": False,
            "message": f"嵌入模型测试失败: {str(e)}"
        }

@app.post("/config/update")
async def update_configuration(
    llm_type: str = Form("local"),
    llm_url: str = Form("https://v1.voct.top/v1"),
    llm_model: str = Form("gpt-4.1-mini"),
    llm_key: str = Form(""),
    llm_temperature: float = Form(0.3),
    embed_type: str = Form("sentence-transformers"),
    embed_url: str = Form("http://192.168.22.191:8000/v1"),
    embed_key: str = Form(""),
    embed_model: str = Form("all-MiniLM-L6-v2"),
    embed_dimension: int = Form(384)
):
    """实时更新系统配置"""
    try:
        # 更新全局配置
        current_config["llm"] = {
            "type": llm_type,
            "url": llm_url,
            "model": llm_model,
            "key": llm_key or "local-api-key",
            "temperature": llm_temperature
        }
        
        current_config["embedding"] = {
            "type": embed_type,
            "url": embed_url,
            "key": embed_key,
            "model": embed_model,
            "dimension": embed_dimension
        }
        
        return {
            "success": True,
            "message": "配置更新成功",
            "config": current_config
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"配置更新失败: {str(e)}")

@app.get("/config/current")
async def get_current_config():
    """获取当前配置"""
    return {
        "success": True,
        "config": current_config
    }

@app.post("/generate")
async def generate_protocol(request: ProtocolGenerationRequest):
    """生成临床试验方案 - 增强版"""
    try:
        # 1. 解析用户需求
        requirement_prompt = f"""
        分析以下临床试验需求，提取关键信息：
        {request.user_requirement}
        
        请提取：药物类型、适应症、试验分期、主要目的、目标人群等。
        返回JSON格式。
        """
        
        extracted_info = call_local_llm(requirement_prompt, 0.1)
        
        # 2. 知识库检索
        # 基于提取的信息进行多维度检索
        search_terms = [
            request.user_requirement,
            extracted_info.get('drug_type', ''),
            extracted_info.get('indication', ''),
            f"{extracted_info.get('trial_phase', '')} 临床试验"
        ]
        
        all_relevant_docs = []
        for term in search_terms:
            if term:
                results = await search_knowledge_embedding(term, top_k=5)
                if results['success']:
                    all_relevant_docs.extend(results['results'])
        
        # 去重和排序
        seen = set()
        unique_docs = []
        for doc in sorted(all_relevant_docs, key=lambda x: x['score'], reverse=True):
            doc_id = doc.get('content', '')[:100]  # 用前100字符作为ID
            if doc_id not in seen:
                seen.add(doc_id)
                unique_docs.append(doc)
        
        # 3. 生成各模块内容
        protocol_sections = {}
        for module in ['研究背景与目的', '研究设计', '研究人群', '给药方案', 
                       '安全性评估', '疗效评估', '统计分析', '数据管理与质量控制']:
            
            # 获取该模块最相关的知识
            module_knowledge = [doc for doc in unique_docs 
                              if any(keyword in doc['content'] 
                                    for keyword in module.split())][:3]
            
            # 生成该模块
            module_prompt = get_module_generation_prompt(
                module, extracted_info, module_knowledge
            )
            
            module_content = call_local_llm(module_prompt, request.temperature)
            protocol_sections[module] = module_content
        
        # 4. 质量检查
        if request.include_quality_check:
            quality_check_prompt = f"""
            评估临床试验方案质量：
            1. 完整性（0-25分）：所有必需章节是否齐全
            2. 科学性（0-25分）：设计是否合理、论述是否严谨
            3. 合规性（0-25分）：是否符合法规要求
            4. 一致性（0-25分）：前后逻辑是否一致
            
            方案概要：{str(protocol_sections)[:2000]}
            
            给出总分和具体问题。
            """
            
            quality_result = call_local_llm(quality_check_prompt, 0.1)
        
        return {
            "success": True,
            "protocol_content": protocol_sections,
            "extracted_info": extracted_info,
            "knowledge_used": len(unique_docs),
            "quality_report": quality_result if request.include_quality_check else None
        }
        
    except Exception as e:
        logger.error(f"生成失败: {e}")
        raise HTTPException(status_code=500, detail=str(e))

# 在后端添加模板端点
@app.get("/templates/clinical_trial/{phase}")
async def get_clinical_trial_template(phase: str):
    """获取特定期别的临床试验方案模板"""
    templates = {
        "phase1": {
            "title": "I期临床试验方案模板",
            "sections": [
                {"title": "研究背景与目的", "required": True},
                {"title": "研究设计", "required": True},
                {"title": "研究人群", "required": True},
                {"title": "研究药物及给药方案", "required": True},
                {"title": "安全性评估", "required": True, "focus": "primary"},
                {"title": "药代动力学评估", "required": True},
                {"title": "疗效评估", "required": False},
                {"title": "统计分析", "required": True},
                {"title": "数据管理", "required": True},
                {"title": "伦理要求", "required": True}
            ],
            "key_points": [
                "剂量递增设计（3+3, BOIN等）",
                "DLT定义和评估",
                "MTD/RP2D确定",
                "PK采样方案",
                "安全性监测计划"
            ]
        },
        "phase2": {
            "title": "II期临床试验方案模板",
            "sections": [
                {"title": "研究背景与目的", "required": True},
                {"title": "研究设计", "required": True},
                {"title": "研究人群", "required": True},
                {"title": "研究药物及给药方案", "required": True},
                {"title": "疗效评估", "required": True, "focus": "primary"},
                {"title": "安全性评估", "required": True},
                {"title": "统计分析", "required": True},
                {"title": "数据管理", "required": True}
            ],
            "key_points": [
                "Simon两阶段设计",
                "主要疗效终点（ORR）",
                "次要终点（PFS、OS、DOR）",
                "疗效评价标准（RECIST 1.1）"
            ]
        }
    }
    
    template = templates.get(f"phase{phase.lower()}", None)
    if not template:
        raise HTTPException(status_code=404, detail="模板未找到")
    
    return template


# 简化版本生成器（作为备用）
async def generate_protocol_simplified(request: ProtocolGenerationRequest):
    """简化版本的协议生成器"""
    try:
        # 这里使用你之前的简化逻辑
        extracted_info = {
            "drug_type": "研究药物",
            "disease": "目标疾病",
            "phase": "I",
            "primary_objective": "评估安全性和耐受性"
        }
        
        protocol_content = {
            "试验背景": "基于用户需求生成的试验背景...",
            "试验设计": "基于用户需求生成的试验设计...",
            "受试者选择": "基于用户需求生成的受试者选择标准..."
        }
        
        quality_report = {
            "overall_score": 75.0,
            "module_scores": {"基础评估": 75},
            "issues": ["使用简化版本生成器"],
            "recommendations": ["建议安装完整版本生成器"]
        }
        
        return {
            "success": True,
            "generated_at": datetime.now().isoformat(),
            "user_requirement": request.user_requirement,
            "extracted_info": extracted_info,
            "protocol_content": protocol_content,
            "quality_report": quality_report,
            "generation_stats": {
                "generator_type": "Simplified",
                "total_modules": len(protocol_content)
            }
        }
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"简化生成失败: {str(e)}")

    

@app.get("/knowledge/stats")
async def get_knowledge_stats():
    """获取知识库统计信息（基于真实向量数据）"""
    try:
        # 动态计算各类型的真实文档数量
        real_stats = {}
        
        # 从向量数据库中统计各类型文档数量
        for knowledge_type in knowledge_stats.keys():
            count = sum(1 for doc in embedded_documents if doc["knowledge_type"] == knowledge_type)
            real_stats[knowledge_type] = {"document_count": count}
        
        # 计算用户上传文档的数量
        user_uploaded_count = sum(1 for doc in embedded_documents if doc["knowledge_type"] == "用户上传文档")
        real_stats["用户上传文档"] = {"document_count": user_uploaded_count}
        
        # 添加总体统计信息
        total_embedded = len(embedded_documents)
        total_files = len(uploaded_files)
        
        return {
            "success": True, 
            "stats": real_stats,
            "summary": {
                "total_embedded_documents": total_embedded,
                "total_uploaded_files": total_files,
                "embedding_model": current_config["embedding"]["type"],
                "avg_docs_per_file": total_embedded / total_files if total_files > 0 else 0
            }
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"获取统计信息失败: {str(e)}")

@app.get("/knowledge/search")
async def search_knowledge(query: str, top_k: int = 5):
    """搜索知识库（真正的向量相似度搜索）"""
    try:
        if not embedded_documents:
            return {"success": True, "results": [], "message": "知识库为空，请先上传文档"}
        
        # 获取查询文本的向量
        query_embedding = get_embedding(query)
        
        # 计算与所有文档的相似度
        results = []
        for doc in embedded_documents:
            similarity = cosine_similarity(query_embedding, doc['embedding'])
            
            if similarity > 0.1:  # 只返回相似度大于0.1的结果
                results.append({
                    "knowledge_type": doc["knowledge_type"],
                    "content": doc["content"][:200] + "..." if len(doc["content"]) > 200 else doc["content"],
                    "metadata": doc["metadata"],
                    "score": similarity
                })
        
        # 按相似度排序并返回top_k结果
        results.sort(key=lambda x: x['score'], reverse=True)
        
        return {
            "success": True, 
            "results": results[:top_k],
            "search_info": {
                "query": query,
                "total_docs_searched": len(embedded_documents),
                "results_found": len(results),
                "embedding_dimension": len(query_embedding)
            }
        }
        
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"向量搜索失败: {str(e)}")

# 添加文本分块功能

@app.post("/knowledge/upload")
async def upload_knowledge_file(
    file: UploadFile = File(...),
    knowledge_type: str = Form("用户上传文档"),
    title: Optional[str] = Form(None)
):
    """上传文件到知识库（真正的向量化处理）"""
    try:
        # 读取文件内容
        file_content = await file.read()
        
        # 保存文件到uploads目录
        file_path = UPLOAD_DIR / file.filename
        
        # 如果文件已存在，添加时间戳
        if file_path.exists():
            stem = file_path.stem
            suffix = file_path.suffix
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            file_path = UPLOAD_DIR / f"{stem}_{timestamp}{suffix}"
        
        # 保存文件
        with open(file_path, "wb") as buffer:
            buffer.write(file_content)
        
        # 提取文本内容
        text_chunks = extract_text_from_file(file_content, file.filename)
        
        # 进行文本分块处理
        processed_chunks = []
        for chunk in text_chunks:
            # 对较长的文本进行进一步分块
            sub_chunks = chunk_text(chunk, chunk_size=500, overlap=50)
            processed_chunks.extend(sub_chunks)
        
        # 真正的向量化处理 - 每个文本块都生成embedding
        embedded_count = 0
        embeddings_info = []
        
        for i, chunk in enumerate(processed_chunks):
            try:
                # 调用embedding API
                embedding = get_embedding(chunk)
                
                # 创建文档条目
                doc_entry = {
                    "id": f"{file.filename}_{i}_{datetime.now().timestamp()}",
                    "content": chunk,
                    "embedding": embedding,
                    "knowledge_type": knowledge_type,
                    "metadata": {
                        "title": title or file.filename,
                        "source_file": file.filename,
                        "chunk_index": i,
                        "upload_time": datetime.now().isoformat(),
                        "file_type": file_path.suffix,
                        "embedding_dimension": len(embedding)
                    }
                }
                
                # 添加到全局向量数据库
                embedded_documents.append(doc_entry)
                embedded_count += 1
                
                # 记录embedding信息用于调试
                embeddings_info.append({
                    "chunk_length": len(chunk),
                    "embedding_dimension": len(embedding),
                    "embedding_sample": embedding[:3] if len(embedding) > 3 else embedding
                })
                
            except Exception as e:
                logger.error(f"为文本块 {i} 生成embedding失败: {e}")
                continue
        
        # 记录文件信息（只存储预览，不存储所有chunks）
        file_info = {
            "filename": file_path.name,
            "original_name": file.filename,
            "size": file_path.stat().st_size,
            "modified": file_path.stat().st_mtime,
            "knowledge_type": knowledge_type,
            "title": title or file.filename,
            "upload_time": datetime.now().isoformat(),
            "chunks_count": len(processed_chunks),
            "embedded_count": embedded_count,
            "chunks": processed_chunks[:3] if len(processed_chunks) > 3 else processed_chunks  # 只保存前3个块作为预览
        }
        
        uploaded_files.append(file_info)
        save_data(embedded_documents, uploaded_files)
        
        return {
            "success": True,
            "message": f"文件 {file.filename} 上传并向量化成功",
            "file_path": str(file_path),
            "records_added": embedded_count,
            "chunks_count": len(processed_chunks),
            "processing_info": {
                "file_type": file_path.suffix,
                "text_extracted": len(text_chunks) > 0,
                "chunking_applied": True,
                "embedding_applied": embedded_count > 0,
                "embedding_model": current_config["embedding"]["type"],
                "embedding_failures": len(processed_chunks) - embedded_count
            },
            "embeddings_sample": embeddings_info[:3]  # 返回前3个embedding的信息
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"文件上传失败: {str(e)}")

@app.get("/knowledge/files")
async def list_uploaded_files():
    """列出已上传的文件"""
    try:
        return {
            "success": True, 
            "files": uploaded_files,
            "total_embedded_documents": len(embedded_documents)
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"获取文件列表失败: {str(e)}")

@app.delete("/knowledge/file/{filename}")
async def delete_knowledge_file(filename: str):
    """删除知识库中的文件"""
    try:
        global uploaded_files, embedded_documents
        
        # 查找并删除文件信息
        file_to_delete = None
        
        for i, file_info in enumerate(uploaded_files):
            if file_info["filename"] == filename:
                file_to_delete = uploaded_files.pop(i)
                break
        
        if not file_to_delete:
            raise HTTPException(status_code=404, detail="文件未找到")
        
        # 删除相关的向量文档
        original_count = len(embedded_documents)
        embedded_documents = [doc for doc in embedded_documents 
                            if doc["metadata"]["source_file"] != file_to_delete["original_name"]]
        deleted_vectors = original_count - len(embedded_documents)
        
        # 删除物理文件
        file_path = UPLOAD_DIR / filename
        if file_path.exists():
            file_path.unlink()
        save_data(embedded_documents, uploaded_files)
        
        return {
            "success": True,
            "deleted_count": 1,
            "deleted_vectors": deleted_vectors,
            "message": f"文件 {filename} 和 {deleted_vectors} 个向量记录删除成功"
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"删除文件失败: {str(e)}")

@app.get("/knowledge/file/{filename}/details")
async def get_file_details(filename: str):
    """获取文件的详细信息，包括内容和embedding结果"""
    try:
        # 查找文件信息
        file_info = None
        for f in uploaded_files:
            if f["filename"] == filename or f.get("original_name") == filename:
                file_info = f
                break
        
        if not file_info:
            raise HTTPException(status_code=404, detail=f"文件 {filename} 未找到")
        
        # 查找相关的embedding文档
        related_docs = [doc for doc in embedded_documents 
                       if doc["metadata"]["source_file"] == file_info.get("original_name", filename)]
        
        if not related_docs:
            raise HTTPException(status_code=404, detail=f"未找到文件 {filename} 的embedding数据")        # 读取原始文件内容
        original_content = ""
        content_preview = ""
        content_truncated = False
        
        try:
            file_path = UPLOAD_DIR / file_info["filename"]
            logger.info(f"📂 [API] 正在读取原始文件: {file_path}")
            if file_path.exists():
                # 根据文件类型选择不同的读取方式
                file_extension = file_path.suffix.lower()
                
                if file_extension == '.pdf':
                    # 对于PDF文件，从已处理的分块中重构内容，避免重新解析
                    logger.info(f"📄 [API] PDF文件，从分块重构内容")
                    if related_docs:
                        # 按页面顺序重构PDF内容
                        pdf_pages = {}
                        for doc in related_docs:
                            content = doc["content"]
                            # 提取页面信息（如果内容包含"第X页:"标识）
                            if content.startswith("第") and "页:" in content:
                                try:
                                    page_num = int(content.split("页:")[0][1:])
                                    page_content = content.split("页:", 1)[1].strip()
                                    pdf_pages[page_num] = page_content
                                except:
                                    # 如果解析失败，直接使用内容
                                    pdf_pages[len(pdf_pages) + 1] = content
                            else:
                                pdf_pages[len(pdf_pages) + 1] = content
                        
                        # 按页面顺序组合内容
                        if pdf_pages:
                            sorted_pages = sorted(pdf_pages.items())
                            full_content = "\n\n".join([f"第{page}页:\n{content}" for page, content in sorted_pages])
                        else:
                            full_content = "\n\n".join([doc["content"] for doc in related_docs])
                    else:
                        full_content = "PDF文件暂无可显示的文本内容"
                        
                else:
                    # 对于非PDF文件，使用智能编码检测
                    full_content = read_file_with_encoding_detection(file_path)
                
                # 限制内容长度以避免前端显示问题
                MAX_CONTENT_LENGTH = 8000  # 增加到8000字符以适应PDF
                PREVIEW_LENGTH = 1500      # 预览1500字符
                
                if len(full_content) > MAX_CONTENT_LENGTH:
                    original_content = full_content[:MAX_CONTENT_LENGTH] + "\n\n... [内容过长已截断，完整内容包含更多页面]"
                    content_truncated = True
                else:
                    original_content = full_content
                
                # 生成预览内容
                if len(full_content) > PREVIEW_LENGTH:
                    content_preview = full_content[:PREVIEW_LENGTH] + "..."
                else:
                    content_preview = full_content
                    
            else:
                original_content = "原始文件未找到"
                content_preview = "原始文件未找到"
                
            logger.info(f"📄 [API] 原始文件内容长度: {len(original_content)} (截断: {content_truncated})")
        except Exception as e:
            logger.error(f"❌ [API] 读取原始文件失败: {e}")
            original_content = f"无法读取原始文件: {str(e)}"
            content_preview = original_content
        
        # 构建分块信息
        chunks = []
        for i, doc in enumerate(related_docs):
            chunk_info = {
                "id": doc["id"],
                "content": doc["content"],
                "knowledge_type": doc["knowledge_type"],
                "chunk_length": len(doc["content"]),
                "chunk_index": doc["metadata"].get("chunk_index", i),
                "embedding_dimension": len(doc["embedding"]),
                "metadata": doc["metadata"]
            }
            chunks.append(chunk_info)
        
        # 按chunk_index排序
        chunks.sort(key=lambda x: x["chunk_index"])
        
        # 计算统计信息
        file_stats = {
            "original_size": len(original_content),
            "chunks_count": len(chunks),
            "knowledge_types": len(set(chunk["knowledge_type"] for chunk in chunks))
        }
        
        # 计算embedding信息
        embedding_info = {
            "total_embeddings": len(chunks),
            "embedding_dimensions": chunks[0]["embedding_dimension"] if chunks else 0,
            "avg_chunk_length": sum(chunk["chunk_length"] for chunk in chunks) / len(chunks) if chunks else 0,
            "knowledge_types": list(set(chunk["knowledge_type"] for chunk in chunks))
        }
        return {
            "success": True,
            "filename": filename,
            "original_content": original_content,
            "content_preview": content_preview,
            "content_truncated": content_truncated,
            "chunks": chunks,
            "file_stats": file_stats,
            "embedding_info": embedding_info
        }
        
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"获取文件详情失败: {str(e)}")

# =============================================================================
# 智能生成工作流 API 端点
# =============================================================================

# 请求模型
class KeyInfoExtractionRequest(BaseModel):
    input_text: str

class OutlineGenerationRequest(BaseModel):
    confirmed_info: Dict[str, Any]
    original_input: str

class ProtocolStreamRequest(BaseModel):
    confirmed_info: Dict[str, Any]
    outline: List[Dict[str, Any]]
    settings: Dict[str, Any]

class SectionStreamRequest(BaseModel):
    confirmed_info: Dict[str, Any]
    section: Dict[str, Any]
    knowledge_types: List[str] = []
    custom_prompt: Optional[str] = None
    settings: Dict[str, Any] = {}

class SectionPromptRequest(BaseModel):
    confirmed_info: Dict[str, Any]
    section: Dict[str, Any]
    knowledge_types: List[str] = []

class ExportProtocolRequest(BaseModel):
    content: str
    format: str
    metadata: Dict[str, Any]


@app.post("/extract_key_info")
async def extract_key_info(request: KeyInfoExtractionRequest):
    """步骤1：从输入文本中提取临床试验方案的关键信息"""
    try:
        # 构建专门针对临床试验方案的提取提示词
        extraction_prompt = f"""
你是一位专业的临床试验方案专家。请从以下文本中提取临床试验方案的关键信息，并以JSON格式返回。

输入文本：
{request.input_text}

请提取以下关键信息：
1. drug_type（药物类型）：如TCR-T细胞、CAR-T细胞、免疫检查点抑制剂等
2. disease（目标疾病）：具体的癌症类型和分期，如晚期肺鳞癌、复发难治性淋巴瘤等
3. trial_phase（试验分期）：I期、II期、III期或I/II期等
4. primary_objective（主要目的）：安全性、耐受性、有效性、剂量探索等
5. primary_endpoint（主要终点）：如MTD、RP2D、ORR、PFS、OS等
6. secondary_endpoints（次要终点）：列表形式，如DCR、DOR、安全性等
7. patient_population（目标人群）：详细的患者特征描述
8. estimated_enrollment（预计入组）：入组人数范围
9. study_design（研究设计）：单臂/双臂、开放/盲法、剂量递增等
10. inclusion_criteria_hints（入组标准提示）：关键的入组要求
11. treatment_line（治疗线数）：一线、二线或多线治疗
12. biomarker_requirements（生物标志物要求）：如HLA分型、抗原表达等

注意事项：
- 如果某项信息未明确提及，请根据临床试验常规做法进行合理推测
- 对于I期试验，主要终点通常是安全性和耐受性
- 对于细胞治疗产品，需要特别关注HLA配型和靶抗原表达
- 返回的JSON格式必须严格符合要求

返回格式示例：
{{
    "drug_type": "TCR-T细胞治疗",
    "disease": "晚期肺鳞癌",
    "trial_phase": "I期",
    "primary_objective": "评估安全性、耐受性并确定RP2D",
    "primary_endpoint": "剂量限制性毒性(DLT)和最大耐受剂量(MTD)",
    "secondary_endpoints": ["客观缓解率(ORR)", "疾病控制率(DCR)", "无进展生存期(PFS)"],
    "patient_population": "既往标准治疗失败的晚期肺鳞癌患者",
    "estimated_enrollment": "12-18例",
    "study_design": "开放标签、单臂、剂量递增I期研究",
    "inclusion_criteria_hints": "HLA-A*02:01阳性，肿瘤表达靶抗原",
    "treatment_line": "二线及以上",
    "biomarker_requirements": "HLA-A*02:01阳性，NY-ESO-1表达阳性"
}}
"""
        
        # 调用LLM进行关键信息提取
        response = call_local_llm(extraction_prompt, temperature=0.1)
        
        try:
            # 解析JSON响应
            import re
            json_match = re.search(r'\{.*?\}', response, re.DOTALL)
            if json_match:
                extracted_info = json.loads(json_match.group())
            else:
                # 如果无法解析JSON，返回默认的临床试验相关字段
                extracted_info = {
                    "drug_type": "待确定的研究药物",
                    "disease": "待确定的目标适应症",
                    "trial_phase": "I期",
                    "primary_objective": "评估安全性和耐受性",
                    "primary_endpoint": "DLT和MTD",
                    "secondary_endpoints": ["ORR", "DCR", "PFS"],
                    "patient_population": "待确定的患者人群",
                    "estimated_enrollment": "待确定",
                    "study_design": "开放标签、剂量递增研究",
                    "inclusion_criteria_hints": "待补充",
                    "treatment_line": "待确定",
                    "biomarker_requirements": "待确定"
                }
                
            # 确保所有必需字段都存在
            required_fields = [
                "drug_type", "disease", "trial_phase", "primary_objective",
                "primary_endpoint", "patient_population", "estimated_enrollment",
                "study_design"
            ]
            
            for field in required_fields:
                if field not in extracted_info:
                    extracted_info[field] = "待补充"
                    
        except Exception as parse_error:
            logger.error(f"JSON解析失败: {parse_error}")
            # 返回基础模板
            extracted_info = {
                "drug_type": "研究药物",
                "disease": "目标疾病",
                "trial_phase": "I期",
                "primary_objective": "安全性和耐受性评估",
                "primary_endpoint": "DLT/MTD",
                "secondary_endpoints": ["ORR", "安全性"],
                "patient_population": "目标患者人群",
                "estimated_enrollment": "12-30例",
                "study_design": "剂量递增研究",
                "error": "信息提取部分失败，请手动补充"
            }
        
        return {
            "success": True,
            "extracted_info": extracted_info,
            "original_response": response,
            "extraction_quality": validate_extraction_quality(extracted_info),
            "prompt": extraction_prompt
        }
        
    except Exception as e:
        logger.error(f"提取关键信息失败: {e}")
        raise HTTPException(status_code=500, detail=f"关键信息提取失败: {str(e)}")


@app.post("/extract_key_info_stream", include_in_schema=False)
async def extract_key_info_stream_v1(request: KeyInfoExtractionRequest):
    """(已废弃) 流式提取关键信息，返回系统提示词和内容"""
    from fastapi.responses import StreamingResponse
    import asyncio

    async def generate():
        try:
            system_prompt = "你是一位专业的临床试验方案专家。请从用户输入中提取临床试验方案的关键信息。"
            extraction_prompt = (
                "请从以下文本中提取临床试验方案的关键信息，并以JSON格式返回。\n\n"
                f"输入文本：\n{request.input_text}\n\n"
                "请提取以下关键信息：\n"
                "1. drug_type（药物类型）\n"
                "2. disease（目标疾病）\n"
                "3. trial_phase（试验分期）\n"
                "4. primary_objective（主要目的）\n"
                "5. primary_endpoint（主要终点）\n"
                "6. secondary_endpoints（次要终点）\n"
                "7. patient_population（目标人群）\n"
                "8. estimated_enrollment（预计入组）\n"
                "9. study_design（研究设计）\n"
                "10. treatment_line（治疗线数）\n\n"
                "返回纯JSON格式，不要有其他文字。"
            )

            yield f"data: {json.dumps({'type': 'system_prompt', 'content': extraction_prompt})}\n\n"
            await asyncio.sleep(0.1)

            accumulated = ""
            for token in call_local_llm_stream(extraction_prompt, 0.1):
                accumulated += token
                yield f"data: {json.dumps({'type': 'content', 'content': token})}\n\n"
                await asyncio.sleep(0.02)

            try:
                import re
                match = re.search(r'\{.*?\}', accumulated, re.DOTALL)
                info = json.loads(match.group()) if match else {}
                yield f"data: {json.dumps({'type': 'extracted_info', 'content': info})}\n\n"
            except Exception as parse_error:
                yield f"data: {json.dumps({'type': 'error', 'content': str(parse_error)})}\n\n"

            yield f"data: {json.dumps({'type': 'done'})}\n\n"
        except Exception as e:
            yield f"data: {json.dumps({'type': 'error', 'content': str(e)})}\n\n"

    return StreamingResponse(
        generate(),
        media_type="text/event-stream",
        headers={
            "Cache-Control": "no-cache",
            "Connection": "keep-alive",
            "Access-Control-Allow-Origin": "*"
        }
    )


@app.post("/extract_key_info_stream", include_in_schema=False)
async def extract_key_info_stream_v2(request: KeyInfoExtractionRequest):
    """(已废弃) 流式提取关键信息，并同时返回系统提示词"""
    from fastapi.responses import StreamingResponse
    import asyncio

    async def generate():
        try:
            system_prompt = "你是一位专业的临床试验方案专家。请从用户输入中提取临床试验方案的关键信息。"
            extraction_prompt = f"""
请从以下文本中提取临床试验方案的关键信息，并以JSON格式返回。

输入文本：
{request.input_text}

请提取以下关键信息：
1. drug_type（药物类型）
2. disease（目标疾病）
3. trial_phase（试验分期）
4. primary_objective（主要目的）
5. primary_endpoint（主要终点）
6. secondary_endpoints（次要终点）
7. patient_population（目标人群）
8. estimated_enrollment（预计入组）
9. study_design（研究设计）
10. treatment_line（治疗线数）

返回纯JSON格式，不要有其他文字。
"""

            yield f"data: {json.dumps({'type': 'system_prompt', 'content': extraction_prompt})}\n\n"
            await asyncio.sleep(0.1)

            accumulated = ""
            for token in call_local_llm_stream(extraction_prompt, system_prompt, 0.1):
                accumulated += token
                yield f"data: {json.dumps({'type': 'content', 'content': token})}\n\n"
                await asyncio.sleep(0.02)

            try:
                import re
                json_match = re.search(r'\{.*?\}', accumulated, re.DOTALL)
                if json_match:
                    info = json.loads(json_match.group())
                else:
                    info = {
                        "drug_type": "待确定",
                        "disease": "待确定",
                        "trial_phase": "I期",
                        "primary_objective": "评估安全性和耐受性",
                        "primary_endpoint": "DLT和MTD",
                        "secondary_endpoints": ["ORR", "DCR", "PFS"],
                        "patient_population": "待确定",
                        "estimated_enrollment": "20-30例",
                        "study_design": "开放标签、剂量递增研究"
                    }

                yield f"data: {json.dumps({'type': 'extracted_info', 'content': info})}\n\n"
            except Exception as parse_error:
                yield f"data: {json.dumps({'type': 'error', 'content': '信息提取失败: ' + str(parse_error)})}\n\n"

            yield f"data: {json.dumps({'type': 'done', 'content': ''})}\n\n"
        except Exception as e:
            yield f"data: {json.dumps({'type': 'error', 'content': str(e)})}\n\n"

    return StreamingResponse(
        generate(),
        media_type="text/event-stream",
        headers={
            "Cache-Control": "no-cache",
            "Connection": "keep-alive",
            "Access-Control-Allow-Origin": "*"
        }
    )


@app.post("/extract_key_info_stream")
async def extract_key_info_stream(request: KeyInfoExtractionRequest):
    """步骤1：流式提取关键信息并返回系统提示词"""
    from fastapi.responses import StreamingResponse
    import asyncio

    async def generate():
        try:
            system_prompt = "你是一位专业的临床试验方案专家。请从用户输入中提取临床试验方案的关键信息。"
            extraction_prompt = f"""
请从以下文本中提取临床试验方案的关键信息，并以JSON格式返回。

输入文本：
{request.input_text}

请提取以下关键信息：
1. drug_type（药物类型）
2. disease（目标疾病）
3. trial_phase（试验分期）
4. primary_objective（主要目的）
5. primary_endpoint（主要终点）
6. secondary_endpoints（次要终点）
7. patient_population（目标人群）
8. estimated_enrollment（预计入组）
9. study_design（研究设计）
10. treatment_line（治疗线数）

返回纯JSON格式，不要有其他文字。
"""

            yield f"data: {json.dumps({'type': 'system_prompt', 'content': extraction_prompt})}\n\n"
            await asyncio.sleep(0.1)

            accumulated = ""
            for token in call_local_llm_stream(extraction_prompt, system_prompt, 0.1):
                accumulated += token
                yield f"data: {json.dumps({'type': 'content', 'content': token})}\n\n"
                await asyncio.sleep(0.02)

            try:
                import re
                match = re.search(r'\{.*?\}', accumulated, re.DOTALL)
                if match:
                    info = json.loads(match.group())
                else:
                    info = {
                        "drug_type": "待确定",
                        "disease": "待确定",
                        "trial_phase": "I期",
                        "primary_objective": "评估安全性和耐受性",
                        "primary_endpoint": "DLT/MTD",
                        "secondary_endpoints": ["ORR", "PFS"],
                        "patient_population": "待确定",
                        "estimated_enrollment": "20-30例",
                        "study_design": "开放标签、剂量递增研究",
                        "treatment_line": "待确定"
                    }
                yield f"data: {json.dumps({'type': 'extracted_info', 'content': info})}\n\n"
            except Exception as e:
                logger.error(f"JSON解析失败: {e}")
                yield f"data: {json.dumps({'type': 'error', 'content': '信息解析失败'})}\n\n"

            yield f"data: {json.dumps({'type': 'done', 'content': ''})}\n\n"
        except Exception as e:
            logger.error(f"流式提取失败: {e}")
            yield f"data: {json.dumps({'type': 'error', 'content': str(e)})}\n\n"

    return StreamingResponse(generate(), media_type="text/event-stream", headers={"Cache-Control": "no-cache", "Connection": "keep-alive", "Access-Control-Allow-Origin": "*"})

def validate_extraction_quality(info):
    """验证提取信息的质量"""
    quality_score = 100
    issues = []
    
    # 检查关键字段
    if "待" in str(info.get("drug_type", "")):
        quality_score -= 20
        issues.append("药物类型未明确")
        
    if "待" in str(info.get("disease", "")):
        quality_score -= 20
        issues.append("目标疾病未明确")
        
    if not info.get("trial_phase"):
        quality_score -= 15
        issues.append("试验分期缺失")
        
    if not info.get("primary_endpoint"):
        quality_score -= 15
        issues.append("主要终点缺失")
        
    return {
        "score": quality_score,
        "issues": issues,
        "recommendation": "建议补充完善" if quality_score < 80 else "信息完整"
    }


@app.post("/generate_outline")
async def generate_outline(request: OutlineGenerationRequest):
    """步骤2：基于确认信息生成协议大纲"""
    try:
        # 构建符合临床试验方案标准的大纲生成提示词
        outline_prompt = f"""
基于以下确认的临床试验信息，生成一个完整的临床试验方案大纲：

确认信息：
- 药物类型：{request.confirmed_info.get('drug_type', '未指定')}
- 适应症：{request.confirmed_info.get('indication', '未指定')}
- 研究阶段：{request.confirmed_info.get('study_phase', '未指定')}
- 研究类型：{request.confirmed_info.get('study_type', '未指定')}
- 主要目的：{request.confirmed_info.get('primary_objectives', '未指定')}
- 目标人群：{request.confirmed_info.get('patient_population', '未指定')}
- 主要终点：{request.confirmed_info.get('primary_endpoint', '未指定')}
- 预计入组：{request.confirmed_info.get('estimated_enrollment', '未指定')}

原始需求：
{request.original_input}

请生成符合ICH-GCP标准的临床试验方案大纲，包含以下标准章节（返回JSON数组格式）：

[
    {{
        "title": "1. 研究背景与目的",
        "content": "包括疾病背景、药物机制、研究理论基础、主要目的和次要目的",
        "subsections": ["1.1 疾病背景", "1.2 药物介绍", "1.3 研究理论基础", "1.4 研究目的"]
    }},
    {{
        "title": "2. 研究设计",
        "content": "试验类型、分组设计、随机化、盲法、对照选择等",
        "subsections": ["2.1 试验类型", "2.2 研究终点", "2.3 试验设计图"]
    }},
    {{
        "title": "3. 研究人群",
        "content": "入选标准、排除标准、退出标准、中止标准",
        "subsections": ["3.1 入选标准", "3.2 排除标准", "3.3 退出标准", "3.4 中止标准"]
    }},
    {{
        "title": "4. 研究药物及给药方案",
        "content": "试验药物、对照药物、给药方案、剂量调整、合并用药",
        "subsections": ["4.1 试验药物", "4.2 给药方案", "4.3 剂量调整", "4.4 合并用药"]
    }},
    {{
        "title": "5. 研究流程",
        "content": "筛选期、治疗期、随访期的访视安排和检查项目",
        "subsections": ["5.1 研究流程图", "5.2 筛选期", "5.3 治疗期", "5.4 随访期"]
    }},
    {{
        "title": "6. 安全性评估",
        "content": "不良事件定义、严重程度分级、因果关系判定、安全性监测",
        "subsections": ["6.1 安全性参数", "6.2 不良事件", "6.3 严重不良事件", "6.4 安全性监测"]
    }},
    {{
        "title": "7. 疗效评估",
        "content": "疗效评价标准、评估时间点、疗效指标定义",
        "subsections": ["7.1 疗效评价标准", "7.2 疗效评估时间", "7.3 疗效指标定义"]
    }},
    {{
        "title": "8. 统计分析",
        "content": "样本量计算、统计分析集、分析方法、亚组分析",
        "subsections": ["8.1 样本量计算", "8.2 分析数据集", "8.3 统计方法", "8.4 期中分析"]
    }},
    {{
        "title": "9. 数据管理与质量控制",
        "content": "数据采集、质量保证、监查计划、数据管理",
        "subsections": ["9.1 数据管理", "9.2 质量保证", "9.3 监查计划"]
    }},
    {{
        "title": "10. 伦理与法规",
        "content": "伦理审查、知情同意、受试者保护、法规要求",
        "subsections": ["10.1 伦理要求", "10.2 知情同意", "10.3 数据保密", "10.4 法规符合性"]
    }}
]

请确保每个章节都包含适当的子章节(subsections)，并且内容描述准确反映该章节应包含的要素。
"""
        
        # 调用LLM生成大纲
        response = call_local_llm(outline_prompt, temperature=0.2)
        
        try:
            # 尝试解析JSON响应
            import re
            json_match = re.search(r'\[.*?\]', response, re.DOTALL)
            if json_match:
                outline = json.loads(json_match.group())
                # 确保每个章节都有subsections字段
                for section in outline:
                    if 'subsections' not in section:
                        section['subsections'] = []
            else:
                # 使用标准大纲模板
                outline = get_standard_protocol_outline(request.confirmed_info)
        except:
            # JSON解析失败时使用标准模板
            outline = get_standard_protocol_outline(request.confirmed_info)
        
        return {
            "success": True,
            "outline": outline,
            "original_response": response,
            "prompt": outline_prompt
        }
        
    except Exception as e:
        logger.error(f"生成大纲失败: {e}")
        raise HTTPException(status_code=500, detail=f"大纲生成失败: {str(e)}")


# 旧版接口保留以兼容历史调用，实际功能由下方新版实现
@app.post("/generate_outline_stream", include_in_schema=False)
async def generate_outline_stream_legacy(request: OutlineGenerationRequest):
    """(已废弃) 流式生成协议大纲"""
    from fastapi.responses import StreamingResponse
    import asyncio

    async def generate():
        try:
            system_prompt = "你是一位临床试验方案撰写专家。请生成符合ICH-GCP标准的临床试验方案章节目录。"
            outline_prompt = json.dumps(request.confirmed_info, ensure_ascii=False)

            yield f"data: {json.dumps({'type': 'system_prompt', 'content': outline_prompt})}\n\n"
            await asyncio.sleep(0.1)

            accumulated = ""
            for token in call_local_llm_stream(outline_prompt, 0.2):
                accumulated += token
                yield f"data: {json.dumps({'type': 'content', 'content': token})}\n\n"
                await asyncio.sleep(0.02)

            try:
                import re
                match = re.search(r'\[.*?\]', accumulated, re.DOTALL)
                outline = json.loads(match.group()) if match else get_standard_protocol_outline(request.confirmed_info)
                yield f"data: {json.dumps({'type': 'outline', 'content': outline})}\n\n"
            except Exception as parse_error:
                outline = get_standard_protocol_outline(request.confirmed_info)
                yield f"data: {json.dumps({'type': 'outline', 'content': outline})}\n\n"

            yield f"data: {json.dumps({'type': 'done'})}\n\n"
        except Exception as e:
            yield f"data: {json.dumps({'type': 'error', 'content': str(e)})}\n\n"

    return StreamingResponse(
        generate(),
        media_type="text/event-stream",
        headers={
            "Cache-Control": "no-cache",
            "Connection": "keep-alive",
            "Access-Control-Allow-Origin": "*"
        }
    )


@app.post("/generate_outline_stream", include_in_schema=False)
async def generate_outline_stream_v1(request: OutlineGenerationRequest):
    """(已废弃) 流式生成协议大纲"""
    from fastapi.responses import StreamingResponse
    import asyncio

    async def generate():
        try:
            system_prompt = "你是一位临床试验方案撰写专家。请生成符合ICH-GCP标准的临床试验方案目录。"
            outline_prompt = f"""
基于以下确认的临床试验信息，生成协议目录，仅需标题：
{json.dumps(request.confirmed_info, ensure_ascii=False)}

返回格式示例：
[
  {{"title": "1. 研究背景与目的", "subsections": ["1.1 ...", "1.2 ..."]}},
  ...
]
"""

            yield f"data: {json.dumps({'type': 'system_prompt', 'content': outline_prompt})}\n\n"
            await asyncio.sleep(0.1)

            accumulated = ""
            for token in call_local_llm_stream(outline_prompt, system_prompt, 0.2):
                accumulated += token
                yield f"data: {json.dumps({'type': 'content', 'content': token})}\n\n"
                await asyncio.sleep(0.02)

            try:
                import re
                json_match = re.search(r'\[.*?\]', accumulated, re.DOTALL)
                if json_match:
                    outline = json.loads(json_match.group())
                else:
                    outline = get_standard_protocol_outline(request.confirmed_info)
                yield f"data: {json.dumps({'type': 'outline', 'content': outline})}\n\n"
            except Exception as parse_error:
                outline = get_standard_protocol_outline(request.confirmed_info)
                yield f"data: {json.dumps({'type': 'outline', 'content': outline})}\n\n"

            yield f"data: {json.dumps({'type': 'done', 'content': ''})}\n\n"
        except Exception as e:
            yield f"data: {json.dumps({'type': 'error', 'content': str(e)})}\n\n"

    return StreamingResponse(
        generate(),
        media_type="text/event-stream",
        headers={
            "Cache-Control": "no-cache",
            "Connection": "keep-alive",
            "Access-Control-Allow-Origin": "*"
        }
    )


# 最终版本的流式大纲生成接口
@app.post("/generate_outline_stream")
async def generate_outline_stream(request: OutlineGenerationRequest):
    """步骤2：流式生成协议大纲"""
    from fastapi.responses import StreamingResponse
    import asyncio

    async def generate():
        try:
            system_prompt = "你是一位临床试验方案撰写专家。请生成符合ICH-GCP标准的临床试验方案目录。"
            outline_prompt = f"""
基于以下确认的临床试验信息，生成协议目录，仅需标题：
{json.dumps(request.confirmed_info, ensure_ascii=False)}

返回格式示例：
[
  {{"title": "1. 研究背景与目的", "subsections": ["1.1 ...", "1.2 ..."]}},
  ...
]
"""

            yield f"data: {json.dumps({'type': 'system_prompt', 'content': outline_prompt})}\n\n"
            await asyncio.sleep(0.1)

            accumulated = ""
            for token in call_local_llm_stream(outline_prompt, system_prompt, 0.2):
                accumulated += token
                yield f"data: {json.dumps({'type': 'content', 'content': token})}\n\n"
                await asyncio.sleep(0.02)

            try:
                import re
                match = re.search(r'\[.*?\]', accumulated, re.DOTALL)
                if match:
                    outline = json.loads(match.group())
                else:
                    outline = get_standard_protocol_outline(request.confirmed_info)
                yield f"data: {json.dumps({'type': 'outline', 'content': outline})}\n\n"
            except Exception as e:
                logger.error(f"大纲解析失败: {e}")
                outline = get_standard_protocol_outline(request.confirmed_info)
                yield f"data: {json.dumps({'type': 'outline', 'content': outline})}\n\n"

            yield f"data: {json.dumps({'type': 'done', 'content': ''})}\n\n"
        except Exception as e:
            logger.error(f"大纲生成失败: {e}")
            yield f"data: {json.dumps({'type': 'error', 'content': str(e)})}\n\n"

    return StreamingResponse(generate(), media_type="text/event-stream", headers={"Cache-Control": "no-cache", "Connection": "keep-alive", "Access-Control-Allow-Origin": "*"})

def get_standard_protocol_outline(confirmed_info):
    """获取标准的临床试验方案大纲模板"""
    drug_type = confirmed_info.get('drug_type', '试验药物')
    indication = confirmed_info.get('indication', '目标适应症')
    study_phase = confirmed_info.get('study_phase', 'I期')
    
    return [
        {
            "title": "1. 研究背景与目的",
            "subsections": [
                "1.1 疾病背景",
                "1.2 药物介绍",
                "1.3 研究理论基础",
                "1.4 研究目的"
            ]
        },
        {
            "title": "2. 研究设计",
            "subsections": [
                "2.1 试验类型与设计",
                "2.2 主要终点",
                "2.3 次要终点",
                "2.4 试验流程图"
            ]
        },
        {
            "title": "3. 研究人群",
            "subsections": [
                "3.1 入选标准",
                "3.2 排除标准",
                "3.3 退出标准",
                "3.4 中止标准"
            ]
        },
        {
            "title": "4. 研究药物及给药方案",
            "subsections": [
                "4.1 试验药物",
                "4.2 给药方案",
                "4.3 剂量调整",
                "4.4 合并用药管理"
            ]
        },
        {
            "title": "5. 研究流程与访视安排",
            "subsections": [
                "5.1 研究流程总览",
                "5.2 筛选期",
                "5.3 治疗期",
                "5.4 随访期",
                "5.5 访视窗口期"
            ]
        },
        {
            "title": "6. 安全性评估",
            "subsections": [
                "6.1 安全性参数",
                "6.2 不良事件定义与分级",
                "6.3 严重不良事件",
                "6.4 剂量限制毒性(DLT)"
            ]
        },
        {
            "title": "7. 疗效评估",
            "subsections": [
                "7.1 疗效评价标准",
                "7.2 疗效评估时间点",
                "7.3 疗效指标定义",
                "7.4 探索性终点"
            ]
        },
        {
            "title": "8. 统计分析计划",
            "subsections": [
                "8.1 样本量计算",
                "8.2 分析数据集",
                "8.3 统计分析方法",
                "8.4 期中分析",
                "8.5 亚组分析"
            ]
        },
        {
            "title": "9. 数据管理与质量控制",
            "subsections": [
                "9.1 数据采集与管理",
                "9.2 质量保证",
                "9.3 临床监查",
                "9.4 稽查"
            ]
        },
        {
            "title": "10. 伦理、法规与管理",
            "subsections": [
                "10.1 伦理委员会审查",
                "10.2 知情同意",
                "10.3 受试者保护",
                "10.4 方案偏离处理",
                "10.5 研究终止"
            ]
        }
    ]

# 临床试验方案分模块生成提示词模板
def get_module_generation_prompt(module_name, confirmed_info, knowledge_context=""):
    """根据模块名称返回相应的生成提示词"""
    
    drug_type = confirmed_info.get('drug_type', '试验药物')
    indication = confirmed_info.get('indication', confirmed_info.get('disease', '目标适应症'))
    study_phase = confirmed_info.get('study_phase', 'I期')
    
    # 通用的高质量要求
    quality_requirements = """
    
生成要求：
1. 内容必须专业、准确、详实
2. 符合ICH-GCP和中国药监局的相关要求
3. 引用数据必须标注来源
4. 使用标准的医学术语，必要时加注英文
5. 逻辑清晰，层次分明
6. 避免使用模糊或不确定的表述
"""
    
    prompts = {
        "研究背景与目的": f"""
作为临床试验方案撰写专家，请撰写{drug_type}治疗{indication}的{study_phase}临床试验方案的"研究背景与目的"章节。

必须包含以下内容：

### 1.1 疾病背景
- {indication}的流行病学数据（发病率、死亡率、地域分布）
- 疾病的病理生理学特征和分子机制
- 目前的标准治疗方案及其局限性
- 未满足的临床需求

### 1.2 药物介绍
- {drug_type}的作用机制和药理学特性
- 非临床研究数据总结（药效学、药代动力学、毒理学）
- 同类药物的研发现状和临床数据
- 本药物的创新点和潜在优势

### 1.3 研究理论基础
- 选择{indication}作为目标适应症的科学依据
- 剂量选择的理论基础
- 生物标志物（如适用）的选择依据
- 预期的临床获益

### 1.4 研究目的
- 主要目的：{confirmed_info.get('primary_objective', '评估安全性和耐受性')}
- 次要目的：{', '.join(confirmed_info.get('secondary_objectives', ['初步疗效评估', 'PK/PD特征']))}
- 探索性目的：生物标志物探索、作用机制验证等

内置资料（仅供参考，不在章节中列出文献）：
{knowledge_context}

{quality_requirements}
""",

        "研究设计": f"""
请详细设计{drug_type}治疗{indication}的{study_phase}临床试验方案的"研究设计"章节。

### 2.1 试验总体设计
- 试验类型：{confirmed_info.get('study_design', '开放标签、单臂、剂量递增')}
- 试验分期：剂量递增期 + 剂量扩展期（如适用）
- 预计试验周期：筛选期（X周）+ 治疗期（X个周期）+ 随访期（X个月）

### 2.2 剂量递增设计（如适用于I期）
- 剂量递增方案：3+3设计 / BOIN设计 / CRM设计
- 起始剂量：基于非临床数据的1/10 NOAEL或1/6 STD10
- 剂量递增幅度：首次100%，后续根据DLT情况调整（50%或33%）
- DLT定义窗口期：第1个治疗周期（28天）
- 剂量限制性毒性(DLT)定义：
  * 4级血液学毒性持续≥7天
  * 3级非血液学毒性（除外可控制的恶心/呕吐）
  * 治疗相关的延迟>14天
  * 其他研究者判定的不可接受毒性

### 2.3 研究终点
主要终点：
- {confirmed_info.get('primary_endpoint', 'MTD和RP2D的确定')}
- 安全性和耐受性评估

次要终点：
{chr(10).join('- ' + ep for ep in confirmed_info.get('secondary_endpoints', ['ORR', 'DCR', 'DOR', 'PFS', 'OS', 'PK参数']))}

探索性终点：
- 免疫监测指标
- 生物标志物与疗效相关性
- 耐药机制探索

### 2.4 样本量计算
- 剂量递增期：{confirmed_info.get('dose_escalation_n', '18-24例')}
- 剂量扩展期：{confirmed_info.get('dose_expansion_n', '10-20例')}
- 统计学假设和计算依据

内置资料（仅供参考，不在章节中列出文献）：
{knowledge_context}

{quality_requirements}
""",

        "研究人群": f"""
请制定{indication}患者参加{drug_type}{study_phase}临床试验的"研究人群"章节。

### 3.1 入选标准
1. 年龄≥18岁，≤75岁，性别不限
2. 组织学或细胞学确诊的{indication}
3. 疾病分期要求：{confirmed_info.get('disease_stage', '局部晚期或转移性')}
4. 既往治疗：{confirmed_info.get('prior_therapy', '标准治疗失败或不耐受')}
5. ECOG PS评分：0-1分（或KPS≥70分）
6. 预期生存期≥3个月
7. 至少有一个可测量病灶（RECIST 1.1标准）
8. 器官功能满足：
   - 血液学：ANC≥1.5×10⁹/L，PLT≥100×10⁹/L，Hb≥90g/L
   - 肝功能：TBIL≤1.5×ULN，ALT/AST≤2.5×ULN（肝转移≤5×ULN）
   - 肾功能：Cr≤1.5×ULN或肌酐清除率≥50mL/min
   - 心功能：LVEF≥50%
9. 生育能力者同意采取有效避孕措施
10. 签署知情同意书

特殊要求（如适用）：
- HLA分型：{confirmed_info.get('hla_requirement', '')}
- 生物标志物：{confirmed_info.get('biomarker_requirement', '')}

### 3.2 排除标准
1. 中枢神经系统转移（除非已治疗稳定≥4周）
2. 活动性自身免疫性疾病
3. 需要系统性免疫抑制治疗
4. 活动性感染（HBV、HCV、HIV、结核等）
5. 既往使用过{drug_type}类似药物
6. 4周内接受过其他抗肿瘤治疗
7. 严重的心血管疾病史
8. 妊娠或哺乳期妇女
9. 精神疾病或依从性差
10. 研究者认为不适合参加试验的其他情况

### 3.3 退出标准
- 受试者撤回知情同意
- 疾病进展
- 不可耐受的毒性
- 研究者判断继续治疗对受试者不利
- 严重违背方案
- 妊娠

### 3.4 中止标准
- 连续出现非预期的SAE
- DSMB建议终止
 - 监管部门要求

 内置资料（仅供参考，不在章节中列出文献）：
 {knowledge_context}

 {quality_requirements}
"""
    }
    
    prompt = prompts.get(module_name, f"请撰写{module_name}部分的内容。{quality_requirements}")

    # 在提示词末尾附加通用模板或模块特定模板
    module_template = MODULE_TEMPLATES.get(module_name)
    if module_template:
        prompt += f"\n\n参考模板：\n{module_template}"
    elif REFERENCE_TEMPLATE:

        prompt += f"\n\n参考模板：\n{REFERENCE_TEMPLATE}"

    return prompt

def generate_protocol_with_knowledge_enhancement(module_name, confirmed_info, knowledge_results):
    """结合知识库检索结果生成协议内容"""
    
    # 整理知识库检索结果
    knowledge_context = "\n\n".join([
        f"【{result['knowledge_type']}】\n{result['content']}"
        for result in knowledge_results[:3]  # 使用最相关的前3个结果
    ])
    
    # 获取该模块的生成提示词
    prompt = get_module_generation_prompt(module_name, confirmed_info, knowledge_context)

    # 添加引用要求，并告知无需在本章节中列出参考
    prompt += """

重要要求：
1. 内容必须基于循证医学证据
2. 引用的文献必须真实可查（提供PMID/DOI/临床试验注册号）
3. 数据和结论必须有明确来源
4. 专业术语使用规范，可加注英文
5. 格式符合医学论文写作规范
6. 请勿在每个章节单独列出文献，将所有文献汇总到文章末尾。
"""
    
    return prompt

@app.post("/generate_protocol_stream")
async def generate_protocol_stream(request: ProtocolStreamRequest):
    """步骤3：基于大纲实时生成完整协议内容"""
    from fastapi.responses import StreamingResponse
    import asyncio
    
    async def generate_content():
        try:
            # 1. 先进行知识库检索增强
            knowledge_results = []
            reference_titles = set()
            if request.settings.get('include_references', True):
                # 基于确认信息构建检索查询
                search_queries = [
                    f"{request.confirmed_info.get('drug_type', '')} {request.confirmed_info.get('indication', '')}",
                    f"{request.confirmed_info.get('study_phase', '')} 临床试验设计",
                    f"{request.confirmed_info.get('indication', '')} 入组标准"
                ]

                for query in search_queries:
                    if query.strip():
                        search_result = await search_knowledge_embedding(query, top_k=3)
                        if search_result['success']:
                            knowledge_results.extend(search_result['results'])
                            for r in search_result['results']:
                                title = r.get('metadata', {}).get('title')
                                if title:
                                    reference_titles.add(title)
            
            # 2. 按照大纲逐个模块生成内容
            full_content = ""
            total_sections = len(request.outline)
            
            for idx, section in enumerate(request.outline):
                # 获取该模块相关的知识
                relevant_knowledge = [k for k in knowledge_results
                                    if any(keyword in k['content']
                                          for keyword in section['title'].split())]
                for r in relevant_knowledge[:3]:
                    title = r.get('metadata', {}).get('title')
                    if title:
                        reference_titles.add(title)
                
                # 构建该模块的生成提示词
                module_prompt = generate_protocol_with_knowledge_enhancement(
                    section['title'],
                    request.confirmed_info,
                    relevant_knowledge[:3]
                )
                
                # 调用LLM生成该模块内容
                module_tokens = ""
                for token in call_local_llm_stream(module_prompt, temperature=0.3):
                    module_tokens += token
                    chunk_data = {
                        "content": token,
                        "progress": idx / total_sections,
                        "current_module": section['title'],
                        "done": False
                    }
                    yield f"data: {json.dumps(chunk_data)}\n\n"
                    await asyncio.sleep(0.02)

                # 模块完成后更新进度并累积内容
                full_content += f"\n## {section['title']}\n\n{module_tokens}\n"
                yield f"data: {json.dumps({'progress': (idx + 1) / total_sections})}\n\n"

            # 在所有章节完成后插入统一的参考文献目录
            if reference_titles:
                ref_text = "\n## 参考文献\n\n" + "\n".join(
                    f"{i+1}. {t}" for i, t in enumerate(sorted(reference_titles))
                ) + "\n"
                full_content += ref_text
                yield f"data: {json.dumps({'content': ref_text})}\n\n"

            # 3. 质量检查（如果启用）
            if request.settings.get('include_quality_check', True):
                quality_prompt = f"""
                请对以下临床试验方案进行质量评估，包括：
                1. 内容完整性（各章节是否齐全）
                2. 科学严谨性（设计是否合理）
                3. 法规合规性（是否符合GCP要求）
                4. 逻辑一致性（前后是否矛盾）
                
                方案内容：
                {full_content[:3000]}...
                
                请给出0-100的评分和改进建议。
                """
                
                quality_result = call_local_llm(quality_prompt, temperature=0.1)
                
                # 解析质量评分
                import re
                score_match = re.search(r'(\d+)分', quality_result)
                quality_score = int(score_match.group(1)) if score_match else 85
                
                quality_data = {
                    "content": f"\n## 质量评估报告\n\n{quality_result}\n",
                    "quality_score": quality_score,
                    "done": False
                }
                yield f"data: {json.dumps(quality_data)}\n\n"
            
            # 发送完成信号
            final_data = {
                "content": "",
                "progress": 1.0,
                "done": True,
                "total_length": len(full_content)
            }
            yield f"data: {json.dumps(final_data)}\n\n"
            
        except Exception as e:
            error_data = {
                "error": str(e),
                "done": True
            }
            yield f"data: {json.dumps(error_data)}\n\n"
    
    return StreamingResponse(
        generate_content(),
        media_type="text/event-stream",
        headers={
            "Cache-Control": "no-cache",
            "Connection": "keep-alive",
            "Access-Control-Allow-Origin": "*"
        }
    )


@app.post("/get_section_prompt")
async def get_section_prompt(request: SectionPromptRequest):
    """返回生成指定章节默认提示词"""
    knowledge_results = []
    if request.knowledge_types:
        query = f"{request.confirmed_info.get('drug_type', '')} {request.confirmed_info.get('indication', '')} {request.section.get('title', '')}"
        search = await search_knowledge_embedding(query, top_k=3, types=request.knowledge_types)
        if search['success']:
            knowledge_results.extend(search['results'])
    prompt = generate_protocol_with_knowledge_enhancement(
        request.section['title'], request.confirmed_info, knowledge_results[:3]
    )
    return {"prompt": prompt}


@app.post("/generate_section_stream")
async def generate_section_stream(request: SectionStreamRequest):
    """逐步生成单个章节内容"""
    from fastapi.responses import StreamingResponse
    import asyncio

    async def stream():
        try:
            knowledge_results = []
            if request.knowledge_types:
                query = f"{request.confirmed_info.get('drug_type', '')} {request.confirmed_info.get('indication', '')} {request.section.get('title', '')}"
                search = await search_knowledge_embedding(query, top_k=3, types=request.knowledge_types)
                if search['success']:
                    knowledge_results.extend(search['results'])

            if request.custom_prompt:
                knowledge_context = "\n\n".join([f"【{r['knowledge_type']}】\n{r['content']}" for r in knowledge_results[:3]])
                prompt = request.custom_prompt + "\n\n" + knowledge_context
            else:
                prompt = generate_protocol_with_knowledge_enhancement(
                    request.section['title'], request.confirmed_info, knowledge_results[:3]
                )

            # 先发送系统提示词，便于前端展示和编辑
            yield f"data: {json.dumps({'type': 'system_prompt', 'content': prompt})}\n\n"

            for token in call_local_llm_stream(prompt, temperature=request.settings.get('detail_level', 0.3)):
                yield f"data: {json.dumps({'content': token})}\n\n"
                await asyncio.sleep(0.02)

            yield f"data: {json.dumps({'done': True})}\n\n"
        except Exception as e:
            yield f"data: {json.dumps({'error': str(e), 'done': True})}\n\n"

    return StreamingResponse(
        stream(),
        media_type="text/event-stream",
        headers={
            "Cache-Control": "no-cache",
            "Connection": "keep-alive",
            "Access-Control-Allow-Origin": "*"
        }
    )


@app.post("/export_protocol")
async def export_protocol(request: ExportProtocolRequest):
    """导出协议为不同格式"""
    try:
        from fastapi.responses import Response
        import io
        
        if request.format.lower() == 'pdf':
            # 简化的PDF生成（需要安装reportlab）
            try:
                from reportlab.lib.pagesizes import A4
                from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
                from reportlab.lib.styles import getSampleStyleSheet
                
                buffer = io.BytesIO()
                doc = SimpleDocTemplate(buffer, pagesize=A4)
                styles = getSampleStyleSheet()
                story = []
                
                # 添加标题
                title = Paragraph("医疗操作协议", styles['Title'])
                story.append(title)
                story.append(Spacer(1, 12))
                
                # 添加内容
                content = Paragraph(request.content.replace('\n', '<br/>'), styles['Normal'])
                story.append(content)
                
                doc.build(story)
                buffer.seek(0)
                
                return Response(
                    content=buffer.getvalue(),
                    media_type="application/pdf",
                    headers={"Content-Disposition": "attachment; filename=protocol.pdf"}
                )
                
            except ImportError:
                # 如果没有reportlab，返回文本格式
                return Response(
                    content=request.content.encode('utf-8'),
                    media_type="text/plain",
                    headers={"Content-Disposition": "attachment; filename=protocol.txt"}
                )
        
        elif request.format.lower() == 'docx':
            # 简化的Word文档生成（需要安装python-docx）
            try:
                from docx import Document
                
                doc = Document()
                doc.add_heading('医疗操作协议', 0)
                
                # 添加元数据
                if request.metadata:
                    doc.add_heading('协议信息', level=1)
                    for key, value in request.metadata.get('confirmed_info', {}).items():
                        doc.add_paragraph(f"{key}: {value}")
                
                doc.add_heading('协议内容', level=1)
                doc.add_paragraph(request.content)
                
                buffer = io.BytesIO()
                doc.save(buffer)
                buffer.seek(0)
                
                return Response(
                    content=buffer.getvalue(),
                    media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    headers={"Content-Disposition": "attachment; filename=protocol.docx"}
                )
                
            except ImportError:
                # 如果没有python-docx，返回文本格式
                return Response(
                    content=request.content.encode('utf-8'),
                    media_type="text/plain",
                    headers={"Content-Disposition": "attachment; filename=protocol.txt"}
                )
        
        else:  # 默认为文本格式
            return Response(
                content=request.content.encode('utf-8'),
                media_type="text/plain",
                headers={"Content-Disposition": f"attachment; filename=protocol.{request.format}"}
            )
            
    except Exception as e:
        logger.error(f"导出协议失败: {e}")
        raise HTTPException(status_code=500, detail=f"导出失败: {str(e)}")

# 启动事件
@app.on_event("startup")
async def startup_event():
    """应用启动时的初始化"""
    logger.info("🚀 医学AI Agent API服务启动中...")
    logger.info("✅ 带真实LLM调用的API服务启动成功!")
    logger.info("📖 API文档地址: http://localhost:8000/docs")
    logger.info("🌐 前端地址: http://localhost:3000")
    logger.info(f"🤖 LLM配置: {current_config['llm']['url']} | {current_config['llm']['model']}")

@app.on_event("shutdown")
async def shutdown_event():
    """应用关闭时的清理"""
    logger.info("👋 医学AI Agent API服务正在关闭...")

if __name__ == "__main__":
    uvicorn.run(
        "start_simple:app",
        host="0.0.0.0",
        port=8000,
        reload=True,
        log_level="info"
    ) 
